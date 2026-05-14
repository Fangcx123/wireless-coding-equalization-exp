"""将 REPORT.md 转换为 REPORT.docx。"""
import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = 'REPORT.md'
DST = 'REPORT.docx'


def add_image_if_exists(doc, alt_text, img_path):
    """如果图片文件存在则插入，否则插入占位文本。"""
    if os.path.exists(img_path):
        try:
            doc.add_picture(img_path, width=Inches(5.0))
            last_paragraph = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption = doc.add_paragraph(f'图：{alt_text}')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].font.size = Pt(10)
            caption.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        except Exception:
            doc.add_paragraph(f'[图片：{alt_text}] ({img_path})')
    else:
        doc.add_paragraph(f'[图片缺失：{alt_text}] ({img_path})')


def set_cell_font(cell, text, bold=False, size=10):
    """设置表格单元格字体。"""
    cell.text = ''
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.bold = bold


def convert():
    with open(SRC, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    list_buffer = []

    def flush_code_block():
        nonlocal code_lines
        if code_lines:
            code_text = ''.join(code_lines).rstrip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(60, 60, 60)
            code_lines = []

    def flush_table():
        if table_rows:
            rows = []
            for tr in table_rows:
                cells = [c.strip() for c in tr.strip('|').split('|')]
                rows.append(cells)
            col_count = max(len(r) for r in rows) if rows else 1
            table = doc.add_table(rows=len(rows), cols=col_count, style='Light Grid Accent 1')
            for i, row_data in enumerate(rows):
                for j, cell_text in enumerate(row_data):
                    if j < col_count:
                        set_cell_font(table.cell(i, j), cell_text, bold=(i == 0), size=10)
            doc.add_paragraph()
            table_rows.clear()

    i = 0
    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.startswith('```'):
            if in_code_block:
                flush_code_block()
                in_code_block = False
            else:
                flush_table()
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 空行
        if line.strip() == '':
            i += 1
            continue

        # 表格行
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if any(c.isalpha() for c in line):  # 跳过纯分隔行如 |---|---|
                in_table = True
                table_rows.append(line.strip())
            i += 1
            continue
        else:
            if in_table:
                flush_table()
                in_table = False

        # 标题
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = re.sub(r'<.*?>', '', heading_match.group(2).strip())
            doc.add_heading(text, level=min(level, 3))
            i += 1
            continue

        # 图片
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            alt = img_match.group(1)
            path = img_match.group(2)
            add_image_if_exists(doc, alt, path)
            i += 1
            continue

        # 无序列表
        list_match = re.match(r'^(\s*)[-*]\s+(.*)', line)
        ordered_match = re.match(r'^(\s*)\d+[\.)]\s+(.*)', line)
        if list_match:
            text = list_match.group(2)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        if ordered_match:
            text = ordered_match.group(2)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # 普通段落：处理行内格式
        text = line.strip()
        text = re.sub(r'<br\s*/?>', '', text)
        text = text.replace('&nbsp;', ' ')

        # 处理 $$ 块级公式 -> 简化为居中文本
        if text == '$$':
            formula_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('$$'):
                formula_lines.append(lines[i].strip())
                i += 1
            i += 1
            if formula_lines:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(' '.join(formula_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(80, 80, 80)
            continue

        # 行内公式 $...$ 和 $$...$$，保留原样
        p = doc.add_paragraph()

        # 简化处理：处理行内 **粗体** 和普通文本
        parts = re.split(r'(\*\*.*?\*\*|`.*?`|\$.*?\$)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            elif part.startswith('$') and part.endswith('$'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.italic = True
            elif part.strip():
                run = p.add_run(part)

        i += 1

    flush_code_block()
    flush_table()

    doc.save(DST)
    print(f'已生成 {DST}')


if __name__ == '__main__':
    convert()
